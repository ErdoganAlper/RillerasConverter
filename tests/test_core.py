"""Tests for the conversion engine.

These run headlessly — nothing here imports Tk.
"""

from __future__ import annotations

import importlib.util
import threading

import fitz
import pytest
from PIL import Image

from rilleras import core


# ------------------------------------------------------------------ helpers --

def make_pdf(path, pages: int = 3, text: str = "Hello"):
    doc = fitz.open()
    for i in range(pages):
        page = doc.new_page()
        page.insert_text((72, 100), f"{text} page {i + 1}", fontsize=24)
    doc.save(str(path))
    doc.close()
    return path


def make_image(path, size=(120, 80), colour=(200, 60, 60)):
    Image.new("RGB", size, colour).save(str(path))
    return path


# ------------------------------------------------------------------- basics --

def test_natural_key_sorts_numerically():
    names = ["page10.png", "page2.png", "page1.png"]
    assert sorted(names, key=core.natural_key) == ["page1.png", "page2.png", "page10.png"]


@pytest.mark.parametrize("text,expected", [
    ("all", [0, 1, 2, 3, 4]),
    ("", [0, 1, 2, 3, 4]),
    ("1", [0]),
    ("1-3", [0, 1, 2]),
    ("1-3,5", [0, 1, 2, 4]),
    ("3-", [2, 3, 4]),
    ("-2", [0, 1]),
    ("99", []),                 # out of range is dropped, not an error
    ("3-1", []),                # reversed range yields nothing
])
def test_parse_page_range(text, expected):
    assert core.parse_page_range(text, 5) == expected


def test_parse_page_range_rejects_garbage():
    with pytest.raises(core.ConversionError):
        core.parse_page_range("abc", 5)


def test_collect_images_sorted_and_filtered(tmp_path):
    make_image(tmp_path / "b10.png")
    make_image(tmp_path / "b2.png")
    (tmp_path / "notes.txt").write_text("ignore me")
    sub = tmp_path / "sub"
    sub.mkdir()
    make_image(sub / "c1.png")

    flat = core.collect_images(tmp_path, recursive=False)
    assert [p.name for p in flat] == ["b2.png", "b10.png"]

    deep = core.collect_images(tmp_path, recursive=True)
    assert len(deep) == 3


# -------------------------------------------------------------------- images --

def test_image_to_image_converts_single_file(tmp_path):
    src = make_image(tmp_path / "in.png")
    out = tmp_path / "out.jpg"
    core.image_to_image(src, out, "jpg", recursive=False, max_size=None, quality=85)

    assert out.exists()
    with Image.open(out) as im:
        assert im.format == "JPEG"


def test_image_to_image_resizes_folder(tmp_path):
    src_dir = tmp_path / "src"
    src_dir.mkdir()
    make_image(src_dir / "a.png", size=(400, 300))
    out_dir = tmp_path / "out"

    core.image_to_image(src_dir, out_dir, "jpg", recursive=False, max_size=100, quality=80)

    result = out_dir / "a.jpg"
    assert result.exists()
    with Image.open(result) as im:
        assert max(im.size) <= 100


def test_rgba_png_to_jpg_flattens_alpha(tmp_path):
    src = tmp_path / "alpha.png"
    Image.new("RGBA", (50, 50), (0, 128, 255, 40)).save(src)
    out = tmp_path / "flat.jpg"

    core.image_to_image(src, out, "jpg", recursive=False, max_size=None, quality=90)

    with Image.open(out) as im:
        assert im.mode == "RGB"


def test_images_to_pdf(tmp_path):
    src = tmp_path / "imgs"
    src.mkdir()
    for i in (1, 2, 10):
        make_image(src / f"p{i}.png")
    out = tmp_path / "album.pdf"

    core.images_to_pdf(src, out, recursive=False, sort_mode="natural")

    with fitz.open(out) as doc:
        assert len(doc) == 3


def test_images_to_pdf_without_images_raises(tmp_path):
    empty = tmp_path / "empty"
    empty.mkdir()
    with pytest.raises(core.ConversionError):
        core.images_to_pdf(empty, tmp_path / "x.pdf", recursive=False, sort_mode="natural")


def test_images_to_pdf_per_subfolder(tmp_path):
    root = tmp_path / "root"
    for name in ("chapter1", "chapter2"):
        d = root / name
        d.mkdir(parents=True)
        make_image(d / "a.png")
    out = tmp_path / "pdfs"

    core.images_to_pdf_per_subfolder(root, out, recursive=False, sort_mode="natural")

    assert (out / "chapter1.pdf").exists()
    assert (out / "chapter2.pdf").exists()


# ----------------------------------------------------------------------- PDF --

def test_pdf_to_images_one_per_page(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=3)
    out = tmp_path / "pages"

    core.pdf_to_images(pdf, out, dpi=72, fmt="png", jpg_quality=90)

    assert sorted(p.name for p in out.glob("*.png")) == ["page_1.png", "page_2.png", "page_3.png"]


def test_pdf_to_images_honours_page_range(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=5)
    out = tmp_path / "pages"

    core.pdf_to_images(pdf, out, dpi=72, fmt="png", jpg_quality=90, page_range="2-3")

    assert sorted(p.name for p in out.glob("*.png")) == ["page_2.png", "page_3.png"]


def test_pdf_to_text(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=2, text="Shadowbound")
    out = tmp_path / "doc.txt"

    core.pdf_to_text(pdf, out)

    content = out.read_text(encoding="utf-8")
    assert "Shadowbound page 1" in content
    assert "Shadowbound page 2" in content


def test_pdf_to_long_image_stitches_pages(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=3)
    out = tmp_path / "long.png"

    core.pdf_to_long_image(pdf, out, dpi=72, fmt="png", jpg_quality=90)

    with fitz.open(pdf) as doc:
        page_height = doc[0].get_pixmap(matrix=fitz.Matrix(1, 1)).height
    with Image.open(out) as im:
        assert im.height >= page_height * 3 - 3  # rounding slack


def test_merge_pdfs(tmp_path):
    a = make_pdf(tmp_path / "a.pdf", pages=2)
    b = make_pdf(tmp_path / "b.pdf", pages=3)
    out = tmp_path / "merged.pdf"

    core.merge_pdfs([a, b], out)

    with fitz.open(out) as doc:
        assert len(doc) == 5


def test_merge_pdfs_empty_list_raises(tmp_path):
    with pytest.raises(core.ConversionError):
        core.merge_pdfs([], tmp_path / "out.pdf")


def test_split_pdf_each_page(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=3)
    out = tmp_path / "parts"

    core.split_pdf(pdf, out, mode="each", ranges="")

    assert len(list(out.glob("*.pdf"))) == 3


def test_split_pdf_by_ranges(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=6)
    out = tmp_path / "parts"

    core.split_pdf(pdf, out, mode="ranges", ranges="1-2,3-6")

    names = sorted(p.name for p in out.glob("*.pdf"))
    assert names == ["doc_1-2.pdf", "doc_3-6.pdf"]
    with fitz.open(out / "doc_3-6.pdf") as doc:
        assert len(doc) == 4


def test_rotate_pdf(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=2)
    out = tmp_path / "rotated.pdf"

    core.rotate_pdf(pdf, out, degrees=90)

    with fitz.open(out) as doc:
        assert doc[0].rotation == 90


def test_compress_pdf_clean_produces_readable_pdf(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=2)
    out = tmp_path / "small.pdf"

    core.compress_pdf_clean(pdf, out)

    with fitz.open(out) as doc:
        assert len(doc) == 2


def test_compress_pdf_rebuild(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=2)
    out = tmp_path / "rebuilt.pdf"

    core.compress_pdf_rebuild(pdf, out, dpi=72)

    with fitz.open(out) as doc:
        assert len(doc) == 2


# ------------------------------------------------------------ cancellation --

def test_cancel_stops_pdf_to_images(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=5)
    out = tmp_path / "pages"
    cancel = threading.Event()
    cancel.set()

    with pytest.raises(core.Cancelled):
        core.pdf_to_images(pdf, out, dpi=72, fmt="png", jpg_quality=90, cancel_event=cancel)


def test_progress_callback_reports_every_page(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=4)
    seen = []

    core.pdf_to_images(pdf, tmp_path / "out", dpi=72, fmt="png", jpg_quality=90,
                       progress_cb=lambda done, total: seen.append((done, total)))

    assert seen == [(1, 4), (2, 4), (3, 4), (4, 4)]


# -------------------------------------------------------------- PDF -> Word --

pdf2docx_installed = importlib.util.find_spec("pdf2docx") is not None


@pytest.mark.skipif(not pdf2docx_installed, reason="pdf2docx not installed")
def test_pdf_to_word_creates_docx(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=2, text="Convertible")
    out = tmp_path / "doc.docx"

    core.pdf_to_word(pdf, out)

    assert out.exists()
    assert out.stat().st_size > 0


@pytest.mark.skipif(not pdf2docx_installed, reason="pdf2docx not installed")
def test_pdf_to_word_forces_docx_extension(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=1)
    core.pdf_to_word(pdf, tmp_path / "wrong.txt")
    assert (tmp_path / "wrong.docx").exists()


@pytest.mark.skipif(not pdf2docx_installed, reason="pdf2docx not installed")
def test_pdf_to_word_respects_page_range(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=4)
    out = tmp_path / "part.docx"

    core.pdf_to_word(pdf, out, page_range="2-3")

    assert out.exists()


def test_pdf_to_word_rejects_empty_selection(tmp_path):
    pdf = make_pdf(tmp_path / "doc.pdf", pages=2)
    with pytest.raises(core.ConversionError):
        core.pdf_to_word(pdf, tmp_path / "out.docx", page_range="50-60")


# -------------------------------------------------------------- Images -> Word --

python_docx_installed = importlib.util.find_spec("docx") is not None
needs_docx = pytest.mark.skipif(not python_docx_installed, reason="python-docx not installed")


@needs_docx
def test_images_to_word_one_picture_per_image(tmp_path):
    from docx import Document

    src = tmp_path / "imgs"
    src.mkdir()
    for i in (1, 2, 10):
        make_image(src / f"p{i}.png")
    out = tmp_path / "album.docx"

    core.images_to_word(src, out, recursive=False, sort_mode="natural")

    assert out.exists()
    assert len(Document(str(out)).inline_shapes) == 3


@needs_docx
def test_images_to_word_scales_wide_image_within_margins(tmp_path):
    from docx import Document

    src = tmp_path / "wide.png"
    make_image(src, size=(4000, 500))
    out = tmp_path / "wide.docx"

    core.images_to_word(src, out, recursive=False, sort_mode="natural")

    doc = Document(str(out))
    section = doc.sections[0]
    usable = section.page_width - section.left_margin - section.right_margin
    shape = doc.inline_shapes[0]
    assert shape.width <= usable


@needs_docx
def test_images_to_word_scales_tall_image_within_page_height(tmp_path):
    from docx import Document

    src = tmp_path / "tall.png"
    make_image(src, size=(500, 4000))
    out = tmp_path / "tall.docx"

    core.images_to_word(src, out, recursive=False, sort_mode="natural")

    doc = Document(str(out))
    section = doc.sections[0]
    usable_h = section.page_height - section.top_margin - section.bottom_margin
    assert doc.inline_shapes[0].height <= usable_h


@needs_docx
def test_images_to_word_handles_webp(tmp_path):
    """python-docx cannot embed WEBP, so it must be converted on the way in."""
    from docx import Document

    src = tmp_path / "shot.webp"
    Image.new("RGB", (200, 150), (40, 90, 200)).save(src, "WEBP")
    out = tmp_path / "shot.docx"

    core.images_to_word(src, out, recursive=False, sort_mode="natural")

    assert len(Document(str(out)).inline_shapes) == 1


@needs_docx
def test_images_to_word_forces_docx_extension(tmp_path):
    src = tmp_path / "a.png"
    make_image(src)
    core.images_to_word(src, tmp_path / "wrong.pdf", recursive=False, sort_mode="natural")
    assert (tmp_path / "wrong.docx").exists()


@needs_docx
def test_images_to_word_without_images_raises(tmp_path):
    empty = tmp_path / "empty"
    empty.mkdir()
    with pytest.raises(core.ConversionError):
        core.images_to_word(empty, tmp_path / "x.docx", recursive=False, sort_mode="natural")


@needs_docx
def test_images_to_word_is_cancellable(tmp_path):
    src = tmp_path / "imgs"
    src.mkdir()
    make_image(src / "a.png")
    cancel = threading.Event()
    cancel.set()

    with pytest.raises(core.Cancelled):
        core.images_to_word(src, tmp_path / "out.docx", recursive=False,
                            sort_mode="natural", cancel_event=cancel)
